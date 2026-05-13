"""
论文 Word 文档工具集
从 3 个内容写作智能体中提取的共享函数。

使用方法：
    import sys, os
    sys.path.insert(0, os.path.expanduser('~/.claude/agents/assets/components'))
    from paper_docx_utils import (
        add_paragraph_with_citations, add_table_from_data,
        load_chapter_plan, collect_references,
        load_numbered_references, add_reference_entry,
        generate_chapter_summary, init_matplotlib_chinese,
    )
"""

import re
import os
import json


# ============================================================
# 引用与段落
# ============================================================

CITE_PATTERN = r'(\[\d+(?:[,，\s]*\d+)*(?:\s*[-–—]\s*\d+)?\])'


def add_paragraph_with_citations(doc, text):
    """添加段落，将 [N] 引用标记转为上标格式。

    引用标记必须作为独立 run 存在，否则上标不会生效。
    """
    from docx import Document  # noqa: ensure docx is importable

    parts = re.split(CITE_PATTERN, text)
    para = doc.add_paragraph()
    for part in parts:
        if not part:
            continue
        if re.match(CITE_PATTERN, part):
            run = para.add_run(part)
            run.font.superscript = True
        else:
            para.add_run(part)
    return para


def add_table_from_data(doc, headers, rows, caption=None):
    """从二维数据创建表格并添加到文档。

    Args:
        doc: python-docx Document 对象
        headers: 表头列表
        rows: 数据行列表（每行是列表）
        caption: 表注文本（可选）
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # 写表头
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = str(h)

    # 写数据
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table.rows[i + 1].cells[j].text = str(cell)

    # 添加表注
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = 1  # 居中

    return table


# ============================================================
# 大纲与文献
# ============================================================

def load_chapter_plan(outline_path, chapter_number):
    """读取指定章节的写作计划。

    Returns:
        dict: {title, subsections, key_points, references, word_target}
    """
    with open(outline_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 简单解析：找到对应章节的 markdown 标题
    # 实际使用时可按 JSON 或 markdown 结构解析
    return {
        "chapter_number": chapter_number,
        "outline_path": outline_path,
    }


def collect_references(chapter_refs_keys, references_path):
    """收集本章需要的文献。

    Args:
        chapter_refs_keys: 本章引用的文献 key 列表
        references_path: references_raw.json 路径

    Returns:
        list: 匹配到的文献列表
    """
    with open(references_path, 'r', encoding='utf-8') as f:
        refs = json.load(f)

    result = []
    for key in chapter_refs_keys:
        for ref in refs:
            if ref.get('key') == key or ref.get('id') == key:
                result.append(ref)
                break
    return result


def load_numbered_references(bib_path, citation_map_path=None):
    """加载参考文献并按编号排列。

    Args:
        bib_path: .bib 文件路径
        citation_map_path: citation_map.md 路径（可选，无则按文件顺序）

    Returns:
        list: 排序后的参考文献列表
    """
    refs = []
    if not os.path.exists(bib_path):
        return refs

    with open(bib_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 简单解析 BibTeX
    entries = re.split(r'\n@', content)
    for entry in entries:
        if not entry.strip():
            continue
        ref = _parse_single_bibtex(entry)
        if ref:
            refs.append(ref)

    # 如果有 citation_map，按引用顺序排序
    if citation_map_path and os.path.exists(citation_map_path):
        with open(citation_map_path, 'r', encoding='utf-8') as f:
            map_content = f.read()
        # 解析 citation_map: key -> number
        order = {}
        for line in map_content.split('\n'):
            m = re.match(r'(\d+)\s*[：:]\s*(\S+)', line.strip())
            if m:
                order[m.group(2)] = int(m.group(1))
        refs.sort(key=lambda r: order.get(r.get('key', ''), 999))

    return refs


def _parse_single_bibtex(entry):
    """解析单条 BibTeX 条目。"""
    m = re.match(r'(\w+)\{([^,]+),', entry.strip())
    if not m:
        return None
    entry_type = m.group(1).lower()
    key = m.group(2).strip()

    fields = {}
    for field_m in re.finditer(r'(\w+)\s*=\s*\{([^}]*)\}', entry):
        fields[field_m.group(1).lower()] = field_m.group(2).strip()

    return {
        'key': key,
        'type': entry_type,
        'authors': fields.get('author', ''),
        'title': fields.get('title', ''),
        'journal': fields.get('journal', fields.get('booktitle', '')),
        'year': fields.get('year', ''),
        'volume': fields.get('volume', ''),
        'pages': fields.get('pages', ''),
    }


def add_reference_entry(doc, ref, number=None):
    """添加单条参考文献条目到文档。

    Args:
        doc: python-docx Document
        ref: 参考文献字典
        number: 编号（如果 ref 中没有 number 字段）
    """
    num = ref.get('number', number)
    entry = f"[{num}] {ref['authors']}. {ref['title']}[{ref['type']}]. {ref['journal']}, {ref['year']}"
    if ref.get('volume'):
        entry += f", {ref['volume']}"
    if ref.get('pages'):
        entry += f": {ref['pages']}"
    entry += "."
    doc.add_paragraph(entry)


# ============================================================
# 章节摘要
# ============================================================

def generate_chapter_summary(chapter_number, core_argument, key_terms,
                             main_evidence, ending_transition, used_citations,
                             output_dir='chapters'):
    """生成章节摘要文件，供下一章智能体参考。

    Args:
        chapter_number: 章节编号
        core_argument: 核心论点（1-2句）
        key_terms: 关键术语 dict {术语: 定义}
        main_evidence: 主要论据列表
        ending_transition: 结尾铺垫文本
        used_citations: 已使用的引用编号列表
        output_dir: 输出目录

    Returns:
        str: 摘要文件路径
    """
    os.makedirs(output_dir, exist_ok=True)
    path = os.path.join(output_dir, f'chapter_{chapter_number}_summary.md')

    with open(path, 'w', encoding='utf-8') as f:
        f.write(f'# 第{chapter_number}章 章节摘要\n\n')
        f.write(f'## 核心论点\n{core_argument}\n\n')
        f.write('## 关键术语\n')
        for term, defn in key_terms.items():
            f.write(f'- {term}：{defn}\n')
        f.write('\n## 主要论据\n')
        for i, ev in enumerate(main_evidence, 1):
            f.write(f'{i}. {ev}\n')
        f.write(f'\n## 结尾铺垫\n{ending_transition}\n\n')
        f.write(f'## 已使用的引用编号\n{", ".join(str(c) for c in used_citations)}\n')

    return path


# ============================================================
# Matplotlib 中文配置
# ============================================================

def init_matplotlib_chinese():
    """初始化 matplotlib 中文字体配置。"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
    plt.rcParams['axes.unicode_minus'] = False
    try:
        plt.style.use('seaborn-v0_8-whitegrid')
    except OSError:
        try:
            plt.style.use('seaborn-whitegrid')
        except OSError:
            pass
    return plt
