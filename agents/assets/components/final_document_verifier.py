"""
final_document_verifier.py — 论文正文一致性验证器

检查最终生成的论文文本/docx 中的数字、日期、模型结论是否来自
results_master.json 或 tables/*.csv。

用法：
    python final_document_verifier.py --project-root .
    python final_document_verifier.py --document outputs/final.docx --results artifacts/results_master.json
    python final_document_verifier.py --document outputs/final.md --results artifacts/results_master.json
"""

import argparse
import csv
import json
import os
import re
import sys
from datetime import datetime
from typing import Any


# ---------------------------------------------------------------------------
# 文档读取
# ---------------------------------------------------------------------------

def read_document(doc_path: str) -> str:
    """读取文档内容，返回纯文本"""
    ext = os.path.splitext(doc_path)[1].lower()

    if ext in (".md", ".txt"):
        with open(doc_path, "r", encoding="utf-8") as f:
            return f.read()

    if ext == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要 python-docx: pip install python-docx")

        doc = Document(doc_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    raise ValueError("不支持的文件格式: {} (支持 .md, .txt, .docx)".format(ext))


# ---------------------------------------------------------------------------
# 结果账本加载
# ---------------------------------------------------------------------------

def load_results_master(path: str) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def load_all_csv_values(project_root: str) -> dict:
    """扫描 tables/*.csv，收集所有数值"""
    csv_values = {}
    tables_dir = os.path.join(project_root, "artifacts", "tables")
    if not os.path.isdir(tables_dir):
        return csv_values

    for fname in os.listdir(tables_dir):
        if not fname.endswith(".csv"):
            continue
        fpath = os.path.join(tables_dir, fname)
        try:
            with open(fpath, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    for cell in row:
                        cell = cell.strip()
                        try:
                            v = float(cell.replace(",", "").replace("%", ""))
                            csv_values[cell] = v
                        except ValueError:
                            pass
        except Exception:
            pass

    return csv_values


def collect_known_numbers(data: dict, csv_values: dict) -> set:
    """从 results_master.json、derived_numbers 和 CSV 收集所有已知数值的字符串表示"""
    known = set()

    def _collect(obj, prefix=""):
        if isinstance(obj, dict):
            for k, v in obj.items():
                _collect(v, "{}.{}".format(prefix, k) if prefix else k)
        elif isinstance(obj, list):
            for item in obj:
                _collect(item, prefix)
        elif isinstance(obj, (int, float)) and not isinstance(obj, bool):
            known.add(str(obj))
            if isinstance(obj, float):
                known.add("{:.4f}".format(obj))
                known.add("{:.2f}".format(obj))
                # 只对 >= 10 的数添加 {.1f}，避免 0.9751 → 1.0 误匹配
                if abs(obj) >= 10:
                    known.add("{:.1f}".format(obj))
                # 百分比兼容：如果值在 (0, 1) 范围内，添加 *100 的表示
                # 例如 0.748 → 添加 "74.8" 以匹配 "74.8%"
                if 0 < abs(obj) < 1:
                    pct = obj * 100
                    known.add("{:.1f}".format(pct))
                    known.add("{:.2f}".format(pct))
                    known.add("{:.0f}".format(pct))
            known.add(str(int(obj)) if isinstance(obj, float) and obj == int(obj) else str(obj))
        elif isinstance(obj, str):
            try:
                v = float(obj.replace(",", "").replace("%", ""))
                known.add(str(v))
                known.add(obj)
            except ValueError:
                pass

    _collect(data)

    # 收集 derived_numbers 中的值
    derived = data.get("derived_numbers", {})
    for key, entry in derived.items():
        v = entry.get("value")
        if v is not None and isinstance(v, (int, float)):
            known.add(str(v))
            if isinstance(v, float):
                known.add("{:.4f}".format(v))
                known.add("{:.2f}".format(v))
                if abs(v) >= 10:
                    known.add("{:.1f}".format(v))
                # 百分比兼容：值在 (0,1) 范围 → 添加 *100 表示
                if 0 < abs(v) < 1:
                    pct = v * 100
                    known.add("{:.1f}".format(pct))
                    known.add("{:.2f}".format(pct))
                    known.add("{:.0f}".format(pct))
            # 百分比兼容：derived_numbers 中 unit="%" 的值存储为整数形式（如 4.5 表示 4.5%），
            # 文本中为 "4.5%"，提取时 /100=0.045，需要反向添加兼容格式
            unit = entry.get("unit", "")
            if unit == "%" and isinstance(v, (int, float)):
                frac = v / 100.0
                known.add("{:.4f}".format(frac))
                known.add("{:.3f}".format(frac))
                known.add("{:.2f}".format(frac))

    for csv_val in csv_values.values():
        known.add(str(csv_val))
        known.add("{:.4f}".format(csv_val))
    return known


# ---------------------------------------------------------------------------
# 内容抽取
# ---------------------------------------------------------------------------

# 关键段落模式：摘要、模型评价、预测结果、优化结果、套餐评价、结论
KEY_SECTION_PATTERNS = [
    r"摘\s*要", r"abstract",
    r"模型.*(?:评价|检验|选择|对比|比较|建立|求解)",
    r"(?:预测|预报).*(?:结果|效果|精度)",
    r"优化.*(?:结果|方案|利润|模型)",
    r"套餐.*(?:设计|评价|营养|达标)",
    r"(?:结论|总结|建议)",
    r"问题[一二三四五].*(?:模型|结果|求解)",
    r"^[一二三四五六七八九十]+[、.].*(?:模型|结果|求解|评价)",
]

# A类数字：关键结果指标（必须在账本中）—— 缺失报 error
A_CLASS_PATTERNS = [
    r"R[²2]", r"RMSE", r"MAE", r"MAPE", r"MSE",
    r"(?:准确|精确|召回)率",
    r"(?:利润|收益|成本|营业额)",
    r"(?:达标|覆盖|满足|合格)率",
    r"(?:提升|增长|增加|提高|减少|降低)(?:了|约|达|幅度)?",
    r"(?:最优|最佳|最好|最差)",
    r"(?:预测|预报).*(?:人数|数量|份)",
    r"(?:浪费|损耗)(?:率|量)",
]

# B类数字：方法参数（来自 method_config）—— 缺失报 warning
B_CLASS_PATTERNS = [
    r"(?:训练|测试).*(?:集|比例|划分|分割)",
    r"(?:交叉验证|折数|k折|K折|K-fold)",
    r"(?:置信|显著).*(?:水平|区间|度)",
    r"(?:学习率|步长|迭代|epoch|batch)",
    r"(?:惩罚|正则).*(?:系数|参数|因子)",
    r"(?:成本|代价).*(?:比|系数|率)",
]

# C类数字：背景信息（年份、页码、序号、数学常数等）—— 直接忽略
C_CLASS_PATTERNS = [
    r"(?:20[12]\d)\s*年",
    r"第\s*\d+\s*[章节页天]",
    r"(?:共|总计|合计)\s*\d+\s*(?:条|项|个|条记录)",
    r"(?:表|图|Figure|Table)\s*\d+",
    r"(?:组合|排列|C\(|P\(|binom|组合数|排列数)",
    r"(?:千万|百万|亿)",
]

DATE_SECTION_PATTERNS = [
    r"问题[三3]",
    r"优化.*日期",
    r"备餐.*日期",
    r"^[一二三四五六七八九十]+[、.].*(?:优化|日期|备餐)",
]

# 数字抽取正则（宽松匹配，后续在 extract_numbers_from_text 中做范围表达式修正）
RE_DECIMAL = re.compile(r"[-+]?\d+\.?\d*")
RE_PERCENT = re.compile(r"(\d+\.?\d*)\s*%")
RE_DATE_ISO = re.compile(r"(\d{4})-(\d{1,2})-(\d{1,2})")
RE_DATE_CN = re.compile(r"(\d{1,2})\s*月\s*(\d{1,2})\s*日")

# 模型名
MODEL_NAMES = {"Prophet", "XGBoost", "LSTM", "GRU", "ARIMA", "SARIMA",
               "Transformer", "CNN", "Linear", "Ridge", "Lasso", "RandomForest",
               "随机森林", "支持向量机", "神经网络"}

# 模型指标关键字
METRIC_KEYWORDS = {"R2", "RMSE", "MAE", "MAPE", "MSE", "准确率", "精确率", "召回率"}

# 结论强度词
STRONG_CONCLUSIONS = [
    (r"(?:最优|最佳|最好|表现最好|效果最好|表现最佳)", "model_selection"),
    (r"(?:最差|最劣|表现最差)", "model_selection"),
    (r"(?:提升|增长|增加|提高)(?:了|约|达)?\s*\d+\.?\d*\s*%", "improvement"),
    (r"(?:达标率|覆盖率|满足率)(?:均|都|为|达)?\s*(?:超过|不低于|高于|达到)?\s*\d+\.?\d*\s*%", "nutrition"),
    (r"(?:利润)(?:提升|增加|增长)(?:了|约|达)?\s*\d+\.?\d*\s*%", "profit"),
]


def extract_document_sections(text: str) -> list:
    """将文档文本按章节拆分，返回 [(section_name, section_text)]"""
    lines = text.split("\n")
    sections = []
    current_name = "preamble"
    current_lines = []

    for line in lines:
        stripped = line.strip()
        # 检测章节标题（Markdown # 或中文编号）
        if (re.match(r"^#{1,3}\s", stripped) or
                re.match(r"^[一二三四五六七八九十]+[、.]", stripped) or
                re.match(r"^\d+\.?\d*\s", stripped)):
            if current_lines:
                sections.append((current_name, "\n".join(current_lines)))
            current_name = stripped.lstrip("#").strip()
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_name, "\n".join(current_lines)))

    return sections


def classify_section(name: str) -> str:
    """判断章节是否为关键段落"""
    name_lower = name.lower()
    for pattern in KEY_SECTION_PATTERNS:
        if re.search(pattern, name, re.IGNORECASE):
            return "key"
    for pattern in DATE_SECTION_PATTERNS:
        if re.search(pattern, name, re.IGNORECASE):
            return "date_key"
    return "background"


def _strip_thousands_separators(text: str) -> str:
    """去除数字中的千分位逗号（如 '205,748.04' → '205748.04'，'75287,520' → '75287520'）"""
    # 循环替换直到没有更多逗号需要去除
    prev = None
    while prev != text:
        prev = text
        text = re.sub(r'(\d),(\d)', r'\1\2', text)
    return text


def extract_numbers_from_text(text: str) -> list:
    """从文本中抽取所有数值，返回 [(value_float, raw_string, start, end)]。

    跳过紧跟 % 的数字（由 extract_percentages_from_text 处理）。
    修正范围表达式误识别（如 "700-900" 中的 -900 → 900）。
    修正千分位逗号（如 "205,748.04" → 205748.04）。
    """
    cleaned = _strip_thousands_separators(text)
    results = []
    for m in RE_DECIMAL.finditer(cleaned):
        # 跳过百分比中的数字（后面紧跟 %）
        end_pos = m.end()
        if end_pos < len(cleaned) and cleaned[end_pos] == "%":
            continue
        raw = m.group()
        try:
            v = float(raw)
        except ValueError:
            continue

        # 修正范围表达式误识别：当负数前面是数字、% 或 -（双破折号）时，
        # 说明是范围分隔符，而非真正的负数。
        # 例如 "700-900" → -900 应为 900；"15%-20%" → -20 应为 20；
        # "523.97--524.82" → -524.82 应为 524.82
        if v < 0 and m.start() > 0:
            prev_char = cleaned[m.start() - 1]
            if prev_char.isdigit() or prev_char == '%' or prev_char == '-':
                # 这是范围表达式的右侧，取绝对值
                v = abs(v)
                raw = raw.lstrip('-').lstrip('+')

        results.append((v, raw, m.start(), m.end()))
    return results


def extract_percentages_from_text(text: str) -> list:
    """抽取所有百分比，返回 [(value_as_fraction, raw_string)]"""
    cleaned = _strip_thousands_separators(text)
    results = []
    for m in RE_PERCENT.finditer(cleaned):
        try:
            v = float(m.group(1)) / 100.0
            results.append((v, m.group(0)))
        except ValueError:
            pass
    return results


def extract_dates_from_text(text: str) -> list:
    """抽取所有日期"""
    dates = []
    for m in RE_DATE_ISO.finditer(text):
        try:
            dt = datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
            dates.append((dt, m.group(0)))
        except ValueError:
            pass
    for m in RE_DATE_CN.finditer(text):
        try:
            dt = datetime(2025, int(m.group(1)), int(m.group(2)))
            dates.append((dt, m.group(0)))
        except ValueError:
            pass
    return dates


def extract_model_conclusions(text: str) -> list:
    """抽取模型相关结论，返回 [(model_name, conclusion_type, raw_sentence)]"""
    results = []
    for pattern, ctype in STRONG_CONCLUSIONS:
        for m in re.finditer(pattern, text):
            # 取包含匹配的完整句子
            start = text.rfind(".", 0, m.start()) + 1
            end = text.find(".", m.end())
            if end == -1:
                end = len(text)
            sentence = text[start:end].strip()

            # 检查句子中是否提到模型名
            found_model = None
            for model in MODEL_NAMES:
                if model.lower() in sentence.lower():
                    found_model = model
                    break

            if found_model or ctype in ("improvement", "profit", "nutrition"):
                results.append((found_model, ctype, sentence))
    return results


# ---------------------------------------------------------------------------
# 日期工作日检查
# ---------------------------------------------------------------------------

def is_weekday(dt: datetime) -> bool:
    return dt.weekday() < 5


# ---------------------------------------------------------------------------
# 近似匹配
# ---------------------------------------------------------------------------

def find_closest_value(target: float, known_numbers: set, tolerance: float = 0.01) -> bool:
    """检查 target 是否与 known_numbers 中某个值近似匹配。

    容差策略:
    - 目标在 (0.1, 10) 范围内（如 R²、RMSE 等小数指标）: 绝对容差 0.001
    - 目标在 (10, 1000) 范围内（如利润、人数）: 绝对容差 1
    - 目标 >= 1000: 相对容差 0.1%
    - 目标 < 0.1（如百分比小数 0.297）: 绝对容差 0.005
    """
    for ns in known_numbers:
        try:
            v = float(ns.replace(",", ""))
            abs_diff = abs(v - target)

            if abs(target) < 0.1:
                # 百分比小数（如 0.297）: 容差 0.005
                if abs_diff <= 0.005:
                    return True
            elif abs(target) < 10:
                # 指标小数（如 R²=0.9322, RMSE=0.9751）: 容差 0.001
                if abs_diff <= 0.001:
                    return True
            elif abs(target) < 1000:
                # 中等数值（如利润、人数）: 容差 1
                if abs_diff <= 1:
                    return True
            else:
                # 大数值: 相对容差 0.1%
                ref = max(abs(v), abs(target), 1e-10)
                if abs_diff / ref <= 0.001:
                    return True
        except ValueError:
            pass
    return False


# ---------------------------------------------------------------------------
# 主检查逻辑
# ---------------------------------------------------------------------------

def verify_document(
    doc_path: str,
    results_path: str,
    project_root: str = ".",
    manifest_path: str = None,
) -> dict:
    """验证文档与结果账本的一致性"""
    # 加载数据
    data = load_results_master(results_path)
    csv_values = load_all_csv_values(project_root)
    known_numbers = collect_known_numbers(data, csv_values)

    # 加载图表账本（可选）
    manifest = None
    if manifest_path and os.path.exists(manifest_path):
        with open(manifest_path, "r", encoding="utf-8") as f:
            manifest = json.load(f)
    else:
        default_manifest = os.path.join(project_root, "artifacts", "artifact_manifest.json")
        if os.path.exists(default_manifest):
            with open(default_manifest, "r", encoding="utf-8") as f:
                manifest = json.load(f)

    # 读取文档
    text = read_document(doc_path)

    errors = []
    warnings = []

    # 拆分章节
    sections = extract_document_sections(text)

    # --- 规则1: 数字可追溯性 ---
    _check_numbers_traceability(text, sections, known_numbers, errors, warnings)

    # --- 规则2: 模型最优结论 ---
    _check_model_conclusions(text, data, errors, warnings)

    # --- 规则3: 日期范围 ---
    _check_date_scope(text, data, errors, warnings)

    # --- 规则4: 摘要强结论 ---
    _check_abstract_strong_claims(text, data, known_numbers, errors, warnings)

    # --- 规则5: 图表引用可追溯 ---
    if manifest:
        _check_figure_references(text, manifest, errors, warnings)

    # --- 规则6: 图表结论在 allowed_claims 范围内 ---
    if manifest:
        _check_figure_allowed_claims(text, manifest, errors, warnings)

    # --- 规则7: 表格数值与 CSV 一致 ---
    if manifest:
        _check_table_csv_consistency(doc_path, manifest, project_root, errors, warnings)

    return {
        "status": "failed" if errors else "passed",
        "errors": errors,
        "warnings": warnings
    }


def _classify_number_context(text_around: str) -> str:
    """根据数字周围的上下文判断其 A/B/C 类别。

    Returns:
        "A": 关键结果指标（必须在账本中）→ 缺失报 error
        "B": 方法参数（来自 method_config）→ 缺失报 warning
        "C": 背景信息（忽略）
    """
    for pattern in C_CLASS_PATTERNS:
        if re.search(pattern, text_around):
            return "C"
    for pattern in A_CLASS_PATTERNS:
        if re.search(pattern, text_around):
            return "A"
    for pattern in B_CLASS_PATTERNS:
        if re.search(pattern, text_around):
            return "B"
    return "A"  # 默认归为 A 类（严格）


def _check_numbers_traceability(text: str, sections: list, known_numbers: set,
                                errors: list, warnings: list):
    """规则1: 文档中的关键数字必须能在账本中找到近似匹配（A/B/C 分类）"""
    for section_name, section_text in sections:
        section_type = classify_section(section_name)

        for value, raw, start, end in extract_numbers_from_text(section_text):
            # 跳过年份、页码等常见无关数字
            if 1900 <= value <= 2100:
                continue
            if value == int(value) and abs(value) < 10:
                continue

            if find_closest_value(value, known_numbers, tolerance=0.01):
                continue

            # 提取数字周围的上下文（前后各100字符）
            ctx_start = max(0, start - 100)
            ctx_end = min(len(section_text), end + 100)
            context = section_text[ctx_start:ctx_end]

            num_class = _classify_number_context(context)

            entry = {
                "type": "untracked_number_{}".format(num_class),
                "location": section_name,
                "text": raw,
                "claim": "数字 {} 在结果账本和CSV中未找到近似匹配".format(raw),
                "expected": "数字应来自 results_master.json 或 tables/*.csv",
                "source": "document"
            }

            if num_class == "C":
                continue  # C 类直接忽略
            elif num_class == "B":
                warnings.append(entry)  # B 类只警告
            else:
                # A 类：关键段落报 error，背景段落报 warning
                if section_type == "key":
                    errors.append(entry)
                else:
                    warnings.append(entry)

        # 百分比特殊处理（统一按 A 类，因为百分比通常是关键结果）
        for frac, raw in extract_percentages_from_text(section_text):
            if not find_closest_value(frac, known_numbers, tolerance=0.001):
                entry = {
                    "type": "untracked_percentage",
                    "location": section_name,
                    "text": raw,
                    "claim": "百分比 {} 在结果账本和CSV中未找到近似匹配".format(raw),
                    "expected": "百分比应来自 results_master.json 或 tables/*.csv",
                    "source": "document"
                }
                if section_type == "key":
                    errors.append(entry)
                else:
                    warnings.append(entry)


def _check_model_conclusions(text: str, data: dict, errors: list, warnings: list):
    """规则2: 模型最优结论必须与 best_model 一致"""
    best_model = data.get("problem2", {}).get("best_model", "")
    if not best_model:
        return

    for model_name, ctype, sentence in extract_model_conclusions(text):
        if ctype != "model_selection" or not model_name:
            continue

        # 判断结论方向
        is_best = any(w in sentence for w in ["最优", "最佳", "最好", "表现最好", "效果最好", "表现最佳"])
        is_worst = any(w in sentence for w in ["最差", "最劣", "表现最差"])

        if is_best and model_name != best_model:
            errors.append({
                "type": "model_conclusion_mismatch",
                "location": "document",
                "text": sentence[:80],
                "claim": "正文声称 '{}' 最优，但 results_master 中 best_model='{}'".format(model_name, best_model),
                "expected": "best_model 应为 '{}'".format(best_model),
                "source": "problem2.best_model"
            })
        elif is_worst and model_name == best_model:
            errors.append({
                "type": "model_conclusion_mismatch",
                "location": "document",
                "text": sentence[:80],
                "claim": "正文声称 '{}' 最差，但它是 results_master 中的 best_model".format(model_name),
                "expected": "best_model 不应被标记为最差",
                "source": "problem2.best_model"
            })


def _check_date_scope(text: str, data: dict, errors: list, warnings: list):
    """规则3: workdays_only 时不能出现周末日期"""
    p3 = data.get("problem3", {})
    date_scope = p3.get("date_scope", "")
    if date_scope != "workdays_only":
        return

    # 只检查问题三相关段落
    sections = extract_document_sections(text)
    for section_name, section_text in sections:
        section_type = classify_section(section_name)
        if section_type != "date_key" and "问题三" not in section_name and "优化" not in section_name:
            continue

        for dt, raw in extract_dates_from_text(section_text):
            if not is_weekday(dt):
                errors.append({
                    "type": "weekend_date_in_workdays",
                    "location": section_name,
                    "text": raw,
                    "claim": "date_scope='workdays_only'，但正文出现周末日期: {}".format(raw),
                    "expected": "问题三相关段落不应包含周六/周日日期",
                    "source": "problem3.date_scope"
                })


def _check_abstract_strong_claims(text: str, data: dict, known_numbers: set,
                                   errors: list, warnings: list):
    """规则4: 摘要中的强结论数字必须来自 summary_claims"""
    # 提取摘要段落
    abstract_text = ""
    in_abstract = False
    for line in text.split("\n"):
        stripped = line.strip()
        if re.match(r"^摘\s*要$", stripped) or re.match(r"^#{1,3}\s+摘\s*要", stripped):
            in_abstract = True
            continue
        if in_abstract:
            # 遇到下一个章节标题时停止
            if (re.match(r"^[一二三四五六七八九十]+[、.]", stripped) or
                    re.match(r"^#{1,3}\s", stripped) and "摘要" not in stripped):
                break
            abstract_text += line + "\n"

    if not abstract_text.strip():
        return

    # 从 summary_claims 收集允许的数字
    allowed_numbers = set()
    allowed_models = set()
    claims = data.get("summary_claims", [])
    for claim in claims:
        if "abstract" not in claim.get("scope", []):
            continue
        for num in claim.get("numbers", []):
            v = num.get("value")
            if v is not None:
                allowed_numbers.add(str(v))
                allowed_numbers.add("{:.4f}".format(float(v)))
                allowed_numbers.add("{:.2f}".format(float(v)))
        for sk in claim.get("source_keys", []):
            parts = sk.split(".")
            if "best_model" in sk:
                val = data
                for p in parts:
                    if isinstance(val, dict):
                        val = val.get(p)
                    else:
                        val = None
                        break
                if val:
                    allowed_models.add(str(val))

    # 也把整个 results_master 中的数字加入 allowed
    allowed_numbers.update(known_numbers)

    # 检查摘要中的数字
    for value, raw, start, end in extract_numbers_from_text(abstract_text):
        if 1900 <= value <= 2100:
            continue
        if value == int(value) and abs(value) < 10:
            continue
        if not find_closest_value(value, allowed_numbers, tolerance=0.01):
            errors.append({
                "type": "abstract_unauthorized_number",
                "location": "abstract",
                "text": raw,
                "claim": "摘要中出现账本外数字: {}".format(raw),
                "expected": "摘要数字必须来自 summary_claims 或 results_master.json",
                "source": "summary_claims"
            })

    # 检查摘要中的百分比
    for frac, raw in extract_percentages_from_text(abstract_text):
        if not find_closest_value(frac, allowed_numbers, tolerance=0.001):
            errors.append({
                "type": "abstract_unauthorized_percentage",
                "location": "abstract",
                "text": raw,
                "claim": "摘要中出现账本外百分比: {}".format(raw),
                "expected": "摘要百分比必须来自 summary_claims 或 results_master.json",
                "source": "summary_claims"
            })

    # 检查摘要中的模型最优结论
    for model_name, ctype, sentence in extract_model_conclusions(abstract_text):
        if ctype == "model_selection" and model_name:
            if allowed_models and model_name not in allowed_models:
                errors.append({
                    "type": "abstract_unauthorized_model_claim",
                    "location": "abstract",
                    "text": sentence[:80],
                    "claim": "摘要中声称 '{}' 最优，但不在 summary_claims 允许的模型中".format(model_name),
                    "expected": "摘要模型结论必须与 best_model 一致",
                    "source": "problem2.best_model"
                })


# ---------------------------------------------------------------------------
# 图表引用检查（规则5-7）
# ---------------------------------------------------------------------------

RE_FIGURE_REF = re.compile(r"(?:图|Fig\.?|Figure)\s*(\d+[-‐‑]?\d*)", re.IGNORECASE)
RE_TABLE_REF = re.compile(r"(?:表|Table)\s*(\d+[-‐‑]?\d*)", re.IGNORECASE)


def _check_figure_references(text: str, manifest: dict, errors: list, warnings: list):
    """规则5: 正文中引用的图表必须在 artifact_manifest.json 中有登记"""
    # Build lookup sets including display_id and aliases
    registered_figures = set()
    figure_aliases = set()
    for fig in manifest.get("figures", []):
        registered_figures.add(fig.get("figure_id", ""))
        if fig.get("display_id"):
            figure_aliases.add(fig["display_id"])
        for alias in fig.get("aliases", []):
            figure_aliases.add(alias)

    registered_tables = set()
    table_aliases = set()
    for t in manifest.get("tables", []):
        registered_tables.add(t.get("table_id", ""))
        if t.get("display_id"):
            table_aliases.add(t["display_id"])
        for alias in t.get("aliases", []):
            table_aliases.add(alias)

    for m in RE_FIGURE_REF.finditer(text):
        ref_id = m.group(1).replace("-", "").replace("‐", "").replace("‑", "")
        ref_raw = m.group(0).strip()
        # 尝试匹配 fig{N}-{M} 格式
        matched = False
        for fid in registered_figures:
            fid_clean = fid.replace("fig", "").replace("-", "").replace("‐", "").replace("‑", "")
            if fid_clean == ref_id or fid == "fig" + ref_id:
                matched = True
                break
        # Also check display_id and aliases (e.g. "图1-1")
        if not matched:
            for alias in figure_aliases:
                alias_clean = alias.replace("-", "").replace("‐", "").replace("‑", "")
                if alias_clean == "图" + ref_id or alias == ref_raw:
                    matched = True
                    break
        if not matched:
            warnings.append({
                "type": "unregistered_figure_ref",
                "location": "document",
                "text": m.group(0),
                "claim": "正文引用了图 {}，但该图未在 artifact_manifest.json 中登记".format(m.group(1)),
                "expected": "所有引用的图必须在 artifact_manifest.json 中有登记",
                "source": "artifact_manifest.json"
            })

    for m in RE_TABLE_REF.finditer(text):
        ref_id = m.group(1).replace("-", "").replace("‐", "").replace("‑", "")
        ref_raw = m.group(0).strip()
        matched = False
        for tid in registered_tables:
            tid_clean = tid.replace("tab", "").replace("-", "").replace("‐", "").replace("‑", "")
            if tid_clean == ref_id or tid == "tab" + ref_id:
                matched = True
                break
        # Also check display_id and aliases (e.g. "表1-1")
        if not matched:
            for alias in table_aliases:
                alias_clean = alias.replace("-", "").replace("‐", "").replace("‑", "")
                if alias_clean == "表" + ref_id or alias == ref_raw:
                    matched = True
                    break
        if not matched:
            warnings.append({
                "type": "unregistered_table_ref",
                "location": "document",
                "text": m.group(0),
                "claim": "正文引用了表 {}，但该表未在 artifact_manifest.json 中登记".format(m.group(1)),
                "expected": "所有引用的表必须在 artifact_manifest.json 中有登记",
                "source": "artifact_manifest.json"
            })


def _check_figure_allowed_claims(text: str, manifest: dict, errors: list, warnings: list):
    """规则6: 图表标题/注释中的结论不得超出 allowed_claims 范围"""
    for fig in manifest.get("figures", []):
        fid = fig.get("figure_id", "")
        claims = fig.get("allowed_claims", [])

        if not claims:
            continue

        # 收集所有 allowed_claims 的结论关键词
        allowed_conclusions = [c.get("conclusion", "") for c in claims]

        # 从 figure_id 提取数字部分，如 "fig2-1" → "2-1" 或 "21"
        fig_num = fid.replace("fig", "").replace("tab", "")
        fig_num_nodash = fig_num.replace("-", "")

        # 构造匹配模式：图2-1 或 图21
        fig_pattern = re.compile(
            r"(?:图|Fig\.?|Figure)\s*(?:{}|{})[^\n]{{0,200}}".format(
                re.escape(fig_num), re.escape(fig_num_nodash)
            ),
            re.IGNORECASE
        )
        for m in fig_pattern.finditer(text):
            sentence = m.group(0)
            # 检查句子中是否有强结论词
            for pattern, ctype in STRONG_CONCLUSIONS:
                if re.search(pattern, sentence):
                    # 检查是否在 allowed_claims 范围内
                    found_allowed = False
                    for ac in allowed_conclusions:
                        if re.search(pattern, ac):
                            found_allowed = True
                            break
                    if not found_allowed:
                        warnings.append({
                            "type": "figure_claim_exceeds_allowed",
                            "location": "document",
                            "text": sentence[:100],
                            "claim": "图 {} 的描述中包含超出 allowed_claims 范围的结论".format(fid),
                            "expected": "图表结论应在 artifact_manifest.json 的 allowed_claims 范围内",
                            "source": "artifact_manifest.json"
                        })


def _check_table_csv_consistency(doc_path: str, manifest: dict, project_root: str,
                                  errors: list, warnings: list):
    """规则7: 正文表格中的数值必须与 CSV 数据一致"""
    text = read_document(doc_path)

    for table in manifest.get("tables", []):
        tid = table.get("table_id", "")
        csv_path = table.get("path", "")
        if not csv_path:
            continue

        abs_csv = os.path.join(project_root, csv_path) if not os.path.isabs(csv_path) else csv_path
        if not os.path.exists(abs_csv):
            continue

        # 读取 CSV 数值
        csv_numbers = set()
        try:
            with open(abs_csv, "r", encoding="utf-8") as f:
                reader = csv.reader(f)
                for row in reader:
                    for cell in row:
                        cell = cell.strip()
                        try:
                            v = float(cell.replace(",", "").replace("%", ""))
                            csv_numbers.add(v)
                            csv_numbers.add("{:.4f}".format(v))
                            csv_numbers.add("{:.2f}".format(v))
                        except ValueError:
                            pass
        except Exception:
            continue

        # 在正文中查找该表的内容（简化：检查表ID附近的内容）
        tab_pattern = re.compile(
            r"(?:表|Table)\s*{}".format(re.escape(tid.replace("tab", "").replace("-", ""))),
            re.IGNORECASE
        )
        for m in tab_pattern.finditer(text):
            # 取匹配后 500 字符
            snippet = text[m.start():m.start() + 500]
            for value, raw, start, end in extract_numbers_from_text(snippet):
                if 1900 <= value <= 2100:
                    continue
                if not find_closest_value(value, csv_numbers, tolerance=0.01):
                    warnings.append({
                        "type": "table_number_mismatch",
                        "location": "document",
                        "text": raw,
                        "claim": "表 {} 附近的数字 {} 与 CSV 数据不一致".format(tid, raw),
                        "expected": "表格数字应与 {} 中的数据一致".format(csv_path),
                        "source": csv_path
                    })


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="论文正文一致性验证器")
    parser.add_argument("--project-root", default=".", help="项目根目录")
    parser.add_argument("--document", default=None, help="论文文件路径 (.docx/.md/.txt)")
    parser.add_argument("--results", default=None, help="results_master.json 路径")
    parser.add_argument("--manifest", default=None, help="artifact_manifest.json 路径")
    args = parser.parse_args()

    project_root = os.path.abspath(args.project_root)

    doc_path = args.document or _find_document(project_root)
    results_path = args.results or os.path.join(project_root, "artifacts", "results_master.json")
    manifest_path = args.manifest or os.path.join(project_root, "artifacts", "artifact_manifest.json")

    if not doc_path:
        print(json.dumps({
            "status": "error",
            "errors": [{"type": "document_not_found",
                        "location": project_root,
                        "claim": "未找到论文文件",
                        "expected": "请通过 --document 指定论文路径",
                        "source": "filesystem"}],
            "warnings": []
        }, ensure_ascii=False, indent=2))
        sys.exit(1)

    report = verify_document(doc_path, results_path, project_root, manifest_path)
    print(json.dumps(report, ensure_ascii=False, indent=2))
    sys.exit(0 if report["status"] == "passed" else 1)


def _find_document(project_root: str) -> str:
    """在项目根目录下查找论文文件"""
    candidates = []
    for root, dirs, files in os.walk(project_root):
        # 跳过隐藏目录和 node_modules
        dirs[:] = [d for d in dirs if not d.startswith(".") and d != "node_modules"]
        for f in files:
            if f.startswith("~$"):
                continue
            ext = os.path.splitext(f)[1].lower()
            if ext in (".docx", ".md", ".txt"):
                full = os.path.join(root, f)
                # 优先选包含 "论文" 或 "paper" 的文件
                if "论文" in f or "paper" in f.lower() or "final" in f.lower():
                    candidates.insert(0, full)
                else:
                    candidates.append(full)
    return candidates[0] if candidates else ""


if __name__ == "__main__":
    main()
