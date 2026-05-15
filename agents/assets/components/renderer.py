"""
renderer.py — Markdown → DOCX 渲染器

读取 chapters/*.md，解析 {{figure:fig_xxx}} / {{table:table_xxx}} 占位符，
从 artifact_manifest.json 查找实际路径，插入图片和表格，
生成目录、页码，输出 output/final.docx。

用法：
    python renderer.py --project-root . --chapters-dir chapters --output output/final.docx
"""

import argparse
import csv
import json
import os
import re
import sys
from typing import Dict, List, Optional, Tuple

# ── 依赖检查 ──────────────────────────────────────────────────

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    Image = None


# ── 常量 ─────────────────────────────────────────────────────

# 章节文件排序规则
CHAPTER_ORDER = [
    "chapter_0_cover",
    "chapter_0_abstract",
    "chapter_1",
    "chapter_2",
    "chapter_3",
    "chapter_4",
    "chapter_5",
    "chapter_6",
    "chapter_7_references",
    "chapter_8_appendix",
    "abstract",
]

# 占位符正则
RE_FIGURE = re.compile(r"\{\{figure:\s*([^}]+?)\s*\}\}")
RE_TABLE = re.compile(r"\{\{table:\s*([^}]+?)\s*\}\}")

# 中文字体（长三角数学建模竞赛格式规范）
FONT_HEADING = "SimHei"       # 黑体 — 标题
FONT_BODY = "SimSun"          # 宋体 — 正文
FONT_TITLE = "STZhongsong"    # 华文中宋 — 论文标题
FONT_FALLBACK = "Microsoft YaHei"

# 字号映射（中国字号 → pt）
SIZE_CHU = Pt(24)       # 初号
SIZE_XIAO_CHU = Pt(18)  # 小二号 — 论文标题
SIZE_SAN = Pt(16)       # 三号 — 一级标题
SIZE_XIAO_SAN = Pt(15)  # 小三号
SIZE_SI = Pt(14)        # 四号 — 二级标题
SIZE_XIAO_SI = Pt(12)   # 小四号 — 正文、三级标题、摘要
SIZE_WU = Pt(10.5)      # 五号 — 图表注释、参考文献
SIZE_XIAO_WU = Pt(9)    # 小五号

# 行距
LINE_SPACING_FIXED = Pt(20)  # 正文固定行距20磅

# 页边距
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)


# ── ArtifactManifest 轻量读取 ────────────────────────────────

class ManifestReader:
    """从 artifact_manifest.json 读取图表路径"""

    def __init__(self, manifest_path: str):
        with open(manifest_path, "r", encoding="utf-8") as f:
            self._data = json.load(f)

    def get_figure_path(self, figure_id: str) -> Optional[str]:
        """查找图表图片路径"""
        for fig in self._data.get("figures", []):
            if fig.get("figure_id") == figure_id:
                return fig.get("image_path")
        return None

    def get_table_path(self, table_id: str) -> Optional[str]:
        """查找表格CSV路径"""
        for tab in self._data.get("tables", []):
            if tab.get("table_id") == table_id:
                return tab.get("path")
        return None

    def get_table_info(self, table_id: str) -> Optional[dict]:
        """查找表格完整信息"""
        for tab in self._data.get("tables", []):
            if tab.get("table_id") == table_id:
                return tab
        return None


# ── 图片预处理 ────────────────────────────────────────────────

def prepare_image(src_path: str, tmp_dir: str) -> str:
    """将透明PNG转为白底RGB，兼容DOCX"""
    if Image is None:
        return src_path
    try:
        img = Image.open(src_path)
        has_alpha = img.mode in ("RGBA", "LA") or "transparency" in img.info
        if not has_alpha:
            return src_path
        stem = os.path.splitext(os.path.basename(src_path))[0]
        out_path = os.path.join(tmp_dir, f"{stem}_rgb.png")
        rgba = img.convert("RGBA")
        white = Image.new("RGBA", rgba.size, (255, 255, 255, 255))
        white.alpha_composite(rgba)
        white.convert("RGB").save(out_path, "PNG", optimize=True)
        return out_path
    except Exception:
        return src_path


# ── Word 文档构建 ─────────────────────────────────────────────

class DocxBuilder:
    """构建 DOCX 文档"""

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._tmp_dir = os.path.join(os.path.dirname(__file__), "_renderer_tmp")
        os.makedirs(self._tmp_dir, exist_ok=True)

    def _setup_page(self):
        """A4 页面，上下2.54cm，左右3.17cm（长三角竞赛格式规范）"""
        section = self.doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT

    def _setup_styles(self):
        """设置默认字体和行距（宋体小四号，固定行距20磅）"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_BODY
        font.size = SIZE_XIAO_SI  # 小四号 = 12pt
        style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
        # 固定行距20磅
        pf = style.paragraph_format
        pf.line_spacing = LINE_SPACING_FIXED
        from docx.enum.text import WD_LINE_SPACING
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def _set_run_font(self, run, font_name=FONT_BODY, size=Pt(12), bold=False):
        """设置 run 的字体"""
        run.font.name = font_name
        run.font.size = size
        run.font.bold = bold
        run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def add_heading(self, text: str, level: int = 1):
        """添加标题（格式规范：H1=黑体三号居中，H2=黑体四号左对齐，H3=黑体小四号左对齐）"""
        heading = self.doc.add_heading(text, level=level)
        # 字号和对齐
        if level == 1:
            font_size = SIZE_SAN  # 三号
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif level == 2:
            font_size = SIZE_SI  # 四号
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            font_size = SIZE_XIAO_SI  # 小四号
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = FONT_HEADING
            run.font.size = font_size
            run.font.bold = True
            run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
        # 行距
        heading.paragraph_format.line_spacing = LINE_SPACING_FIXED
        from docx.enum.text import WD_LINE_SPACING
        heading.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        return heading

    def add_paragraph(self, text: str, bold=False, align=None, font_size=None):
        """添加段落（默认宋体小四号，固定行距20磅）"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        if font_size is None:
            font_size = SIZE_XIAO_SI  # 小四号 = 12pt
        self._set_run_font(run, FONT_BODY, font_size, bold)
        if align:
            p.alignment = align
        # 固定行距20磅
        from docx.enum.text import WD_LINE_SPACING
        p.paragraph_format.line_spacing = LINE_SPACING_FIXED
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        return p

    def add_figure(self, image_path: str, caption: str, width_inches: float = 5.5):
        """插入图片（图注宋体五号居中）"""
        if not os.path.exists(image_path):
            self.add_paragraph(f"[图片未找到: {image_path}]", bold=True)
            return

        image_path = prepare_image(image_path, self._tmp_dir)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))

        # 图注：宋体五号居中
        cap_p = self.doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        self._set_run_font(cap_run, FONT_BODY, SIZE_WU)

    def add_table_from_csv(self, csv_path: str, caption: str, max_rows: int = 30):
        """从 CSV 插入表格（表题黑体五号居中，表头黑体五号，数据宋体五号）"""
        if not os.path.exists(csv_path):
            self.add_paragraph(f"[表格未找到: {csv_path}]", bold=True)
            return

        with open(csv_path, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            rows = list(reader)

        if not rows:
            return

        # 表题：黑体五号居中加粗（放在表格上方）
        cap_p = self.doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        self._set_run_font(cap_run, FONT_HEADING, SIZE_WU, bold=True)

        # 限制行数
        display_rows = rows[:max_rows + 1]  # +1 for header
        headers = display_rows[0]
        data_rows = display_rows[1:]

        table = self.doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头：黑体五号居中加粗
        for j, header in enumerate(headers):
            cell = table.rows[0].cells[j]
            cell.text = str(header)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = SIZE_WU
                    run.font.name = FONT_HEADING
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_HEADING)
            self._set_cell_shading(cell, 'D9E2F3')

        # 数据行：宋体五号居中
        for i, row in enumerate(data_rows):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = str(val)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = SIZE_WU
                        run.font.name = FONT_BODY
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)

        if len(rows) > max_rows + 1:
            note = self.add_paragraph(
                f"（注：共 {len(rows)-1} 行数据，仅显示前 {max_rows} 行）",
                font_size=Pt(9)
            )
            note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _set_cell_shading(self, cell, color: str):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def add_page_number(self):
        """添加页脚页码"""
        section = self.doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 域代码: PAGE
        run = p.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run.element.append(fld_char_begin)

        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' PAGE '
        run2.element.append(instr)

        run3 = p.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run3.element.append(fld_char_end)

    def add_toc(self):
        """添加目录域代码（Word 中需按 F9 更新）"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run()
        fld_char_begin = OxmlElement('w:fldChar')
        fld_char_begin.set(qn('w:fldCharType'), 'begin')
        run.element.append(fld_char_begin)

        run2 = p.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2.element.append(instr)

        run3 = p.add_run()
        fld_char_separate = OxmlElement('w:fldChar')
        fld_char_separate.set(qn('w:fldCharType'), 'separate')
        run3.element.append(fld_char_separate)

        run4 = p.add_run('[请在 Word 中右键此处 → 更新域 → 更新整个目录]')
        run4.font.color.rgb = RGBColor(128, 128, 128)
        run4.font.size = Pt(10)

        run5 = p.add_run()
        fld_char_end = OxmlElement('w:fldChar')
        fld_char_end.set(qn('w:fldCharType'), 'end')
        run5.element.append(fld_char_end)

    def save(self, path: str):
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        self.doc.save(path)
        # 清理临时文件
        import shutil
        if os.path.isdir(self._tmp_dir):
            shutil.rmtree(self._tmp_dir, ignore_errors=True)


# ── Markdown 解析与渲染 ──────────────────────────────────────

def parse_markdown_file(filepath: str) -> List[dict]:
    """
    解析 markdown 文件为结构化块列表。
    每个块: {"type": "heading"|"paragraph"|"figure"|"table"|"blank", ...}
    """
    with open(filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()

    blocks = []
    for line in lines:
        line = line.rstrip("\n")

        # 空行
        if not line.strip():
            blocks.append({"type": "blank"})
            continue

        # 标题
        heading_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append({"type": "heading", "level": level, "text": heading_match.group(2).strip()})
            continue

        # 图片占位符
        fig_match = RE_FIGURE.search(line)
        if fig_match:
            fig_id = fig_match.group(1).strip()
            # 提取占位符之外的文字作为图注
            caption = RE_FIGURE.sub("", line).strip()
            blocks.append({"type": "figure", "figure_id": fig_id, "caption": caption})
            continue

        # 表格占位符
        tab_match = RE_TABLE.search(line)
        if tab_match:
            tab_id = tab_match.group(1).strip()
            caption = RE_TABLE.sub("", line).strip()
            blocks.append({"type": "table", "table_id": tab_id, "caption": caption})
            continue

        # 普通段落
        blocks.append({"type": "paragraph", "text": line})

    return blocks


def render_chapter(builder: DocxBuilder, blocks: List[dict], manifest: ManifestReader,
                   fig_dir: str, table_dir: str):
    """将解析后的块渲染到文档"""
    for block in blocks:
        btype = block["type"]

        if btype == "blank":
            continue

        elif btype == "heading":
            builder.add_heading(block["text"], level=block["level"])

        elif btype == "paragraph":
            builder.add_paragraph(block["text"])

        elif btype == "figure":
            fig_id = block["figure_id"]
            fig_path = manifest.get_figure_path(fig_id)
            if fig_path:
                # 尝试相对路径和绝对路径
                if not os.path.isabs(fig_path):
                    fig_path = os.path.join(fig_dir, os.path.basename(fig_path))
                if not os.path.exists(fig_path):
                    fig_path = os.path.join(fig_dir, f"{fig_id}.png")
            else:
                # manifest 没找到，直接尝试 figures/ 目录
                fig_path = os.path.join(fig_dir, f"{fig_id}.png")
            caption = block.get("caption", fig_id)
            builder.add_figure(fig_path, caption)

        elif btype == "table":
            tab_id = block["table_id"]
            tab_path = manifest.get_table_path(tab_id)
            if tab_path:
                if not os.path.isabs(tab_path):
                    tab_path = os.path.join(table_dir, os.path.basename(tab_path))
                if not os.path.exists(tab_path):
                    tab_path = os.path.join(table_dir, f"{tab_id}.csv")
            else:
                # manifest 没找到，直接尝试 tables/ 目录
                tab_path = os.path.join(table_dir, f"{tab_id}.csv")
            caption = block.get("caption", tab_id)
            builder.add_table_from_csv(tab_path, caption)


def sort_chapter_files(files: List[str]) -> List[str]:
    """按预定义顺序排序章节文件"""
    def sort_key(filepath):
        basename = os.path.splitext(os.path.basename(filepath))[0]
        for i, name in enumerate(CHAPTER_ORDER):
            if basename == name:
                return (0, i)
        # chapter_N 格式
        m = re.match(r"chapter_(\d+)", basename)
        if m:
            return (0, int(m.group(1)))
        # 其他文件排最后
        return (1, 0)

    return sorted(files, key=sort_key)


# ── 主入口 ────────────────────────────────────────────────────

def render(project_root: str, chapters_dir: str = "chapters",
           output_path: str = "output/final.docx",
           manifest_path: str = None,
           fig_dir: str = None, table_dir: str = None,
           with_toc: bool = True, with_page_numbers: bool = True):
    """
    主渲染函数。

    Args:
        project_root: 项目根目录
        chapters_dir: 章节目录（相对于 project_root）
        output_path: 输出文件路径（相对于 project_root）
        manifest_path: artifact_manifest.json 路径
        fig_dir: 图片目录
        table_dir: 表格目录
        with_toc: 是否添加目录
        with_page_numbers: 是否添加页码
    """
    # 解析路径
    chapters_abs = os.path.join(project_root, chapters_dir)
    output_abs = os.path.join(project_root, output_path)

    if manifest_path is None:
        manifest_path = os.path.join(project_root, "artifacts", "artifact_manifest.json")
    if fig_dir is None:
        fig_dir = os.path.join(project_root, "artifacts", "figures")
    if table_dir is None:
        table_dir = os.path.join(project_root, "artifacts", "tables")

    # 检查 manifest
    if not os.path.exists(manifest_path):
        print(f"ERROR: artifact_manifest.json not found at {manifest_path}")
        print("  Run stage 7 (artifact_manifest_builder) first.")
        return False

    manifest = ManifestReader(manifest_path)

    # 收集章节文件
    if not os.path.isdir(chapters_abs):
        print(f"ERROR: chapters directory not found: {chapters_abs}")
        return False

    all_md = [f for f in os.listdir(chapters_abs) if f.endswith(".md")]

    # 过滤：排除 summary, requests, expansion_log 等非正文文件
    exclude_suffixes = ("_summary.md", "_requests.md", "_expansion_log.md")
    content_md = [f for f in all_md if not any(f.endswith(s) for s in exclude_suffixes)]

    # 优先使用 _expanded.md 版本：如果 chapter_N_expanded.md 存在，跳过 chapter_N.md
    expanded_bases = {f.replace("_expanded.md", "") for f in content_md if f.endswith("_expanded.md")}
    filtered_md = []
    for f in content_md:
        if f.endswith("_expanded.md"):
            filtered_md.append(f)
        elif f.replace(".md", "") in expanded_bases:
            continue  # 跳过原始版本，expanded 优先
        else:
            filtered_md.append(f)

    md_files = [
        os.path.join(chapters_abs, f)
        for f in filtered_md
    ]
    md_files = sort_chapter_files(md_files)

    if not md_files:
        print(f"ERROR: No .md files found in {chapters_abs}")
        return False

    print(f"  Found {len(md_files)} chapter files:")
    for f in md_files:
        print(f"    - {os.path.basename(f)}")

    # 构建文档
    builder = DocxBuilder()

    for md_file in md_files:
        print(f"  Rendering: {os.path.basename(md_file)}")
        blocks = parse_markdown_file(md_file)
        render_chapter(builder, blocks, manifest, fig_dir, table_dir)
        # 章节之间加分页
        builder.doc.add_page_break()

    # 目录（插入在文档开头，但 python-docx 不支持真正插入到开头，
    # 所以在末尾添加目录域，用户在 Word 中手动移动）
    if with_toc:
        print("  Adding TOC field...")
        builder.add_toc()

    # 页码
    if with_page_numbers:
        print("  Adding page numbers...")
        builder.add_page_number()

    # 保存
    builder.save(output_abs)
    print(f"  Output: {output_abs}")
    print(f"  File size: {os.path.getsize(output_abs) / 1024:.1f} KB")
    return True


# ── CLI ───────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Markdown → DOCX renderer")
    parser.add_argument("--project-root", default=".", help="项目根目录")
    parser.add_argument("--chapters-dir", default="chapters", help="章节目录")
    parser.add_argument("--output", default="output/final.docx", help="输出路径")
    parser.add_argument("--manifest", default=None, help="artifact_manifest.json 路径")
    parser.add_argument("--fig-dir", default=None, help="图片目录")
    parser.add_argument("--table-dir", default=None, help="表格目录")
    parser.add_argument("--no-toc", action="store_true", help="不添加目录")
    parser.add_argument("--no-page-numbers", action="store_true", help="不添加页码")
    args = parser.parse_args()

    print("=" * 60)
    print("  renderer.py — Markdown → DOCX")
    print("=" * 60)

    ok = render(
        project_root=os.path.abspath(args.project_root),
        chapters_dir=args.chapters_dir,
        output_path=args.output,
        manifest_path=args.manifest,
        fig_dir=args.fig_dir,
        table_dir=args.table_dir,
        with_toc=not args.no_toc,
        with_page_numbers=not args.no_page_numbers,
    )

    if ok:
        print("\n  Rendering complete.")
    else:
        print("\n  Rendering FAILED.")
        sys.exit(1)


if __name__ == "__main__":
    main()
